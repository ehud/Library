#!/usr/bin/env python
import urllib2
import re
import ast

from docx import Document
from docx.shared import Inches
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def send_email(SUBJECT='',TEXT=''):
            import smtplib

            gmail_user = "user.name@gmail.com"
            gmail_pwd = "password"
            FROM = 'user.name@gmail.com'
            TO = ['user.name@gmail.com'] #must be a list
            
            # Prepare actual message
            message = """\From: %s\nTo: %s\nSubject: %s\n\n%s
            """ % (FROM, ", ".join(TO), SUBJECT, TEXT)
            try:
                #server = smtplib.SMTP(SERVER)
                server = smtplib.SMTP("smtp.gmail.com", 587) #or port 465 doesn't seem to work!
                server.ehlo()
                server.starttls()
                server.login(gmail_user, gmail_pwd)
                server.sendmail(FROM, TO, message)
                #server.quit()
                server.close()
                # print 'successfully sent the mail'
            except:
                print "failed to send mail"

fp=open("/path/to/isbns_to_order.txt")
books=[]
failed=[]
for line in fp:
    isbn = line.rstrip('\n').split()[0]
    if len(line.rstrip('\n').split(' ',1))>1:
       note=line.rstrip('\n').split(' ',1)[1]
    else:
       note=''
    f = urllib2.urlopen("http://xisbn.worldcat.org/webservices/xid/isbn/"+isbn+"?method=getMetadata&fl=*&format=python")
    res=f.read()
    if ast.literal_eval(res)['stat']=='ok':
       details=ast.literal_eval(res)['list'][0]
       details['note']=note
       books.append(details)
    else:
       print isbn  
       failed.append(line)
fp.close()

document = Document()

section = document.sections[0]
section.page_width, section.page_height = (Inches(8.27),Inches(11.69))
section.left_margin, section.right_margin = (Inches(0.3), Inches(0.3))

document.add_heading('Book Purchase (' + datetime.datetime.now().strftime("%B,%Y")+')', 0)

table = document.add_table(rows=1, cols=9)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Author'
hdr_cells[1].text = 'Title'
hdr_cells[2].text = 'ISBN'
hdr_cells[3].text = 'Publisher+Year'
hdr_cells[4].text = 'Edition'
hdr_cells[5].text = 'Priority'
hdr_cells[6].text = 'Copies'
hdr_cells[7].text = 'Notes'
hdr_cells[8].text = 'Number'
for item in books:
    #print item
    row_cells = table.add_row().cells
    row_cells[0].text = item.get('author','')
    row_cells[1].text = item.get('title','')
    row_cells[2].text = item['isbn'][0]
    row_cells[3].text = item.get('publisher','')+'('+item.get('year','')+')'
    row_cells[4].text = item.get('ed','')
    row_cells[5].text = ''
    row_cells[6].text = ''
    row_cells[7].text = item['note']
    row_cells[8].text = ''

#document.add_page_break()

document.save('/path/to/purchase-'+ datetime.datetime.now().strftime("%B%Y")+'.docx')

print "cleaning order list"
fp=open("/path/to/isbns_to_order.txt","w")
for line in failed:
     print line
     fp.write(line)
fp.close()

# add purchased books to watch list
fp=open("/path/to/watched_isbns.txt","a")
fp.writelines([item['isbn'][0]+'\n' for item in books])
 